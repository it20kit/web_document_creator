from abc import ABC, abstractmethod
import re
from docx import Document
from document_creator.document_block_handlers.tools.document_style_manager import DocumentStyleManager
from document_creator.halpers.document_structure_manager import DocumentStructureManager


class AbstractBlockCreator(ABC):
    """Описывает структуру конструкторов блоков и содержит общие методы конструкторов"""

    def __init__(self):
        self.document_structure_manager = DocumentStructureManager()
        self.style_manager = DocumentStyleManager()

    @abstractmethod
    def create_block_based_on_data(self, document: Document, data: dict) -> None:
        """Создает блок документа по переданным данным"""
        pass

    def add_content_in_paragraph(self, paragraph, text: str, style):
        """Создать заголовок по переданным параметрам"""
        text = paragraph.add_run(text)
        self.style_manager.set_style_for_paragraph(paragraph, style)
        self.style_manager.set_text_style_by_parameters(text, style)

    def are_there_empty_values_in_content(self, content: list) -> bool:
        """Проверяет есть ли в массиве с контентом пустые строки"""
        for i in range(len(content)):
            if self.is_value_empty(content[i]):
                return True
        return False

    @staticmethod
    def create_empty_paragraph(document: Document):
        """Создать пустой абзац, который можно потом использовать для создания абзацев и заголовков"""
        return document.add_paragraph("")

    @staticmethod
    def create_empty_numbered_list_paragraph(document: Document):
        """Создать нумерованный абзац, который можно потом использовать для создания абзацев и заголовков"""
        return document.add_paragraph("", style="List Number")

    @staticmethod
    def is_value_empty(value: str) -> bool:
        """Проверяет пустое ли значение по шаблону"""
        return re.match('{{.+}}', value) is not None

    @staticmethod
    def get_value_from_content(content):
        return content.get('value')
