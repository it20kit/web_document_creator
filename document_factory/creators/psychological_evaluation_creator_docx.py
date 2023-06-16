from docx.enum.text import *
from docx.shared import *
from document_factory.creators.abstract_document_creator import AbstractDocumentCreator
from docx import Document
from document_factory.document_descriptor import DocumentDescriptor
from forms.form_for_psychological_evaluation import FormForPsychologicalEvaluation
from document_factory.tools.document_tools import DocumentTools
from document_factory.document_block_handlers.block_handler_manager import BlockHandlerManager
from document_factory.document_block_handlers.heading_creator import HeadingCreator


class PsychologicalEvaluationCreatorDocx(AbstractDocumentCreator):
    """Создает docx документ с психологическим заключением ребёнка, по выбранным параметрам"""

    def __init__(self):
        self.document_tolls = DocumentTools()

    def create_document(self, form: dict) -> DocumentDescriptor:
        """Создать документ психологической оценки и обернуть его в дескриптор"""
        document = Document()
        self.__set_standard_for_all_document(document)
        for block in form['blocks']:
            self.__add_block(document, block)

        # self.__add_block_one(document, data)
        # self.__add_block_two(document, data)
        # self.__add_block_thee(document, data)
        # self.__add_block_with_conclusion(document, data)
        # self.document_tolls.add_empty_string(document)
        # self.__add_table_with_signatures(document)
        return DocumentDescriptor(self.create_name_document(form), document)

    @staticmethod
    def __add_block(document: Document, block: dict):
        block_type = block.get('type')
        block_creator = BlockHandlerManager().get_block_creators_by_type(block_type)
        block_creator(document, block)

    def create_name_document(self, form: dict) -> str:
        """Создать и вернуть имя документа"""
        name_child = form.get('document_name')
        return name_child

    def save(self, document, name):
        """Сохранить документ"""
        document.save(name)

    def __add_block_one(self, document, data: FormForPsychologicalEvaluation) -> None:
        """Создать и добавить первый блок с контентом в документ"""
        self.__create_block_one(document, data)

    def __create_block_one(self, document: Document, data: FormForPsychologicalEvaluation) -> None:
        """Создать первый блок с контентом"""
        indent = 0.1
        if data.means_of_communication is not None:
            title = "Средство коммуникации: "
            content = data.means_of_communication
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_child_contact is not None:
            title = "Особенности контакта ребёнка: "
            content = data.features_of_child_contact
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.emotional_reaction is not None:
            title = "Эмоциональная реакция на ситуацию обследования: "
            content = data.emotional_reaction
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.reaction_to_approval is not None:
            title = "Реакция на одобрение: "
            content = data.reaction_to_approval
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.reaction_to_remark is not None:
            title = "Реакция на замечание: "
            content = data.reaction_to_remark
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.reaction_to_failure is not None:
            title = "Реакция на неудачу: "
            content = data.reaction_to_failure
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.emotional_state is not None:
            title = "Эмоциональное состояние во время выполнения заданий: "
            content = data.emotional_state
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.communication is not None:
            title = "Общение: "
            content = data.communication
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.reaction_to_the_result is not None:
            title = "Реакция на результат: "
            content = data.reaction_to_the_result
            self.__create_content_for_block(document, title, content,indent=indent)

    def __add_block_two(self, document: Document, data: FormForPsychologicalEvaluation) -> None:
        """Создать и добавить второй блок с контентом в документ"""
        self.__create_heading_by_block_two(document)
        self.__create_block_two(document, data)

    def __create_heading_by_block_two(self, document: Document) -> None:
        """Создать заголовок для второго блока"""
        heading = document.add_paragraph("", style="List Number")
        text = heading.add_run("Особенности деятельности во время обследования")
        font_size = 12
        self.document_tolls.set_text_style_by_parameters(text, font_size, bold=True)

    def __create_block_two(self, document: Document, data: FormForPsychologicalEvaluation) -> None:
        """Создать второй блок с контентом"""
        indent = 0.1
        if data.presence_and_persistence is not None:
            title = "Наличие и стойкость интереса к заданию: "
            content = data.presence_and_persistence
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.understanding_instructions is not None:
            title = "Понимание инструкции: "
            content = data.understanding_instructions
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.indicative_activity is not None:
            title = "Ориентировочная деятельность: "
            content = data.indicative_activity
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.independence_of_completing_tasks is not None:
            title = "Самостоятельность выполнения заданий: "
            content = data.independence_of_completing_tasks
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.nature_of_activity is not None:
            title = "Характер деятельности (целенаправленность и активность): "
            content = data.nature_of_activity
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.pace_and_dynamics_of_activity is not None:
            title = "Темп и динамика деятельности: "
            content = data.pace_and_dynamics_of_activity
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.efficiency is not None:
            title = "Работоспособность: "
            content = data.efficiency
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_regulation_of_activity is not None:
            title = "Особенности регуляции деятельности: "
            content = data.features_of_regulation_of_activity
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.organization_of_assistance is not None:
            title = "Организация помощи: "
            content = data.organization_of_assistance
            self.__create_content_for_block(document, title, content, indent=indent)

    def __add_block_thee(self, document: Document, data: FormForPsychologicalEvaluation) -> None:
        """Создать и добавить третий блок с контентом в документ"""
        self.__create_heading_block_three(document)
        self.__create_block_three(document, data)

    def __create_heading_block_three(self, document: Document) -> None:
        """Создать заголовок третьего блока"""
        font_size_heading = 12
        bold = True

        heading = document.add_paragraph("", style="List Number")
        text_heading = heading.add_run("Особенности познавательной сферы и моторной функции рук")
        self.document_tolls.set_text_style_by_parameters(text_heading, font_size_heading, bold=bold)

    def __create_block_three(self, document: Document,  data: FormForPsychologicalEvaluation) -> None:
        """Создать третий блок с контентом"""
        indent = 0.1
        if data.features_of_attention is not None:
            title = "Особенности внимания: "
            content = data.features_of_attention
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_perception is not None:
            title = "Особенности восприятия: "
            content = data.features_of_perception
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.memory_features is not None:
            title = "Особенности памяти: "
            content = data.memory_features
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_thinking is not None:
            title = "Особенности мышления: "
            content = data.features_of_thinking
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_speech is not None:
            title = "Особенности речи: "
            content = data.features_of_speech
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_imagination is not None:
            title = "Особенности воображения: "
            content = data.features_of_imagination
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_motor_function is not None:
            title = "Особенности моторной функции рук: "
            content = data.features_of_motor_function
            self.__create_content_for_block(document, title, content, indent=indent)

        if data.features_of_large_motor_skills is not None:
            title = "Особенности крупной моторики: "
            content = data.features_of_large_motor_skills
            self.__create_content_for_block(document, title, content, indent=indent)

    def __add_block_with_conclusion(self, document: Document, data: FormForPsychologicalEvaluation) -> None:
        """Создать и добавить блок с заключением"""
        self.__create_conclusion_heading(document)
        self.__create_table_with_conclusion(document, data)

    def __create_conclusion_heading(self, document: Document) -> None:
        """Добавить заголовок к блоку с заключением"""
        font_size = 14
        bold = True
        heading = document.add_paragraph("")
        heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        text_heading = heading.add_run("Выводы и рекомендации по результатам психологического обследования")
        self.document_tolls.set_text_style_by_parameters(text_heading, font_size, bold=bold)

    def __create_table_with_conclusion(self, document: Document, data: FormForPsychologicalEvaluation) -> None:
        """Создать и добавить контент в заключительную таблицу"""
        rows = 5
        cols = 2
        table = self.document_tolls.create_table_by_parameters(document, rows, cols)
        table.columns[0].width = Cm(6)
        table.columns[1].width = Cm(12)
        self.__add_data_to_conclusion_table(table, data)

    def __add_data_to_conclusion_table(self, table, data: FormForPsychologicalEvaluation) -> None:
        """Добавить контент в заключительную таблицу"""
        cell = table.cell(0, 0)
        text = "Заключение:"
        text_size = 12
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
        if data.conclusion is not None:
            cell = table.cell(0, 1)
            text = data.conclusion
            self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

        cell = table.cell(1, 0)
        text = "Прогноз:"
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
        if data.forecast is not None:
            cell = table.cell(1, 1)
            text = data.forecast
            self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

        cell = table.cell(2, 0)
        text = "Выводы о динамике:"
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
        if data.conclusions_about_dynamics is not None:
            cell = table.cell(2, 1)
            text = data.conclusions_about_dynamics
            self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

        cell = table.cell(3, 0)
        text = "Общие рекомендации по сопровождению ребёнка:"
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
        if data.general_recommendations is not None:
            cell = table.cell(3, 1)
            text = data.general_recommendations
            self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

        cell = table.cell(4, 0)
        text = "Рекомендации родителям:"
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT, bold=True)
        if data.recommendations_to_parents is not None:
            cell = table.cell(4, 1)
            text = data.recommendations_to_parents
            self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def __add_table_with_signatures(self, document: Document) -> None:
        """Создать и добавить таблицы с подписями"""
        self.__create_first_table_with_signature(document)
        self.document_tolls.add_empty_string(document)
        self.__create_second_table_with_signature(document)
        self.document_tolls.add_empty_string(document)
        self.__create_third_table_with_signature(document)

    def __create_first_table_with_signature(self, document: Document) -> None:
        """Создать первую таблицу с подписями"""
        rows = 2
        cols = 4
        table = self.document_tolls.create_table_by_parameters(document, rows, cols)
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(1)
        table.columns[2].width = Cm(1)
        table.columns[3].width = Cm(11)

        cell = table.cell(0, 0)
        text = "Педагог-психолог"
        text_size = 12
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

        cell = table.cell(1, 1)
        text = "(подпись)"
        text_size = 8
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.CENTER)

        cell = table.cell(1, 3)
        text = "(Фамилия, инициалы)"
        text_size = 8
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.CENTER)

    def __create_second_table_with_signature(self, document: Document) -> None:
        """Создать вторую таблицу с подписями"""
        rows = 2
        cols = 3
        table = self.document_tolls.create_table_by_parameters(document, rows, cols)
        table.columns[0].width = Cm(6)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(10)

        cell = table.cell(0, 0)
        text = "С заключением ознакомлен (а)"
        text_size = 12
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

        cell = table.cell(1, 1)
        text = "(подпись)"
        text_size = 8
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.CENTER)

        cell = table.cell(1, 2)
        text = "(Фамилия, инициалы)"
        text_size = 8
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.CENTER)

    def __create_third_table_with_signature(self, document: Document) -> None:
        """Создать третью таблицу с подписями"""
        rows = 2
        cols = 3
        table = self.document_tolls.create_table_by_parameters(document, rows, cols)
        table.columns[0].width = Cm(8)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(8)

        cell = table.cell(0, 0)
        text = "С заключением и рекомендациями согласен (на)"
        text_size = 12
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.LEFT)

        cell = table.cell(1, 1)
        text = "(подпись)"
        text_size = 8
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.CENTER)

        cell = table.cell(1, 2)
        text = "(Фамилия, инициалы)"
        text_size = 8
        self.document_tolls.add_content_for_cell_table(cell, text, text_size, WD_PARAGRAPH_ALIGNMENT.CENTER)

    def __create_content_for_block(self, document: Document, title: str, content: str, indent=None) -> None:
        """Создать контент для блоков по параметрам"""
        paragraph = document.add_paragraph("")
        if indent is not None:
            indent_size = Inches(indent)
            paragraph.paragraph_format.left_indent = indent_size
        title = paragraph.add_run(title)
        self.document_tolls.set_text_style_by_parameters(title, 12, WD_PARAGRAPH_ALIGNMENT.LEFT, underline=True)
        content = paragraph.add_run(content)
        self.document_tolls.set_text_style_by_parameters(content, 12, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def __set_standard_for_all_document(self, document: Document) -> None:
        """Установить стандарты для всего документа(размер шрифта, отступы и т.д)"""
        font_size = 12
        font_name = 'Times New Roman'
        bottom = 15
        top = 15
        right = 10
        left = 20
        self.document_tolls.set_style_standard_for_all_document(document, font_size, font_name)
        self.document_tolls.set_padding_sizes_for_all_document(document, bottom, top, right, left)
