from abc import ABC, abstractmethod
import re
from docx import Document
from document_factory.document_block_handlers.tools.document_style_manager import DocumentStyleManager


class AbstractBlockCreator(ABC):
    """Описывает структуру конструкторов блоков и содержит общие методы конструкторов"""

    StyleManager = DocumentStyleManager

    @abstractmethod
    def create_block_based_on_data(self, document: Document, data: dict) -> None:
        """Создает блок документа по переданным данным"""
        pass

    @staticmethod
    def add_content_in_paragraph(paragraph, text: str, style):
        """Создать заголовок по переданным параметрам"""
        text = paragraph.add_run(text)
        AbstractBlockCreator.StyleManager.set_style_for_paragraph(paragraph, style)
        AbstractBlockCreator.StyleManager.set_text_style_by_parameters(text, style)

    def are_there_empty_values_in_content(self, content: list) -> bool:
        """Проверяет есть ли в массиве с контентом пустые строки"""
        for i in range(len(content)):
            if self.is_value_empty(content[i]):
                return True
        return False

    @staticmethod
    def create_empty_paragraph(document: Document):
        """Создать пустой абзац, который можно потом использовать для создания абзацев и заголовков"""
        return document.add_paragraph("")

    @staticmethod
    def create_empty_numbered_list_paragraph(document: Document):
        """Создать нумерованный абзац, который можно потом использовать для создания абзацев и заголовков"""
        return document.add_paragraph("", style="List Number")

    @staticmethod
    def get_content_block_from_data(data: dict) -> list | str:
        """Получить содержимое для создания блока документа из данных"""
        return data.get('content')

    @staticmethod
    def get_style_block_from_data(data: dict) -> dict:
        """Получить стили для создания блока документа из данных"""
        return data.get('parameters').get('style')

    @staticmethod
    def is_value_empty(value: str) -> bool:
        """Проверяет пустое ли значение по шаблону"""
        return re.match('{{.+}}', value) is not None
