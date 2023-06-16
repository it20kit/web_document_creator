from document_factory.document_block_handlers.text_creator import TextCreator
from document_factory.document_block_handlers.paragraph_creator import ParagraphCreator
from docx.shared import *


class TableCreator:
    """Умеет создавать различные таблицы"""

    @staticmethod
    def create_empty_table_by_parameters(document, rows: int, cols: int):
        """Создает пустую таблицу по числу строк и столбцов"""
        table = document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.width = 1
        return table

    def create_table_with_two_columns_by_text_style(
            self, table, data: list, style_one: dict = None, style_two: dict = None
    ):
        for i in range(len(data)):
            translation = data[i].get('translation')
            value = data[i].get('value')
            cell = table.cell(i, 0)
            self.__add_content_for_cell_table(cell, translation, style_one)
            if value is not None:
                cell = table.cell(i, 1)
                self.__add_content_for_cell_table(cell, value, style_two)
        return table

    def create_table_two_columns(self, document, data: dict):
        style_for_first_column = data.get('parameters').get('style').get('style_for_first_column')
        style_for_second_column = data.get('parameters').get('style').get('style_for_second_column')
        rows = data.get('parameters').get('table_size').get('number_of_rows_and_columns').get('rows')
        table = self.create_table_with_two_columns_by_text_style(
            self.create_empty_table_by_parameters(document, rows=rows, cols=2),
            data.get('content'),
            style_for_first_column,
            style_for_second_column
        )
        size_by_table = data.get('parameters').get('table_size').get('size_columns')
        if size_by_table is not None:
            self.set_size_of_table(table, size_by_table)

    @staticmethod
    def __add_content_for_cell_table(cell, content, style: dict) -> None:
        """Добавить текст в ячейку таблицы с переданными параметрами"""
        text = cell.paragraphs[0].add_run(content)
        ParagraphCreator().set_style_for_paragraph(cell.paragraphs[0], style)
        TextCreator().set_text_style_by_parameters(text, style)

    @staticmethod
    def set_size_of_table(table, sizes: list):
        for i in range(len(sizes)):
            table.columns[i].width = Cm(sizes[i])
