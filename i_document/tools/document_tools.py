from docx import Document
from docx.enum.text import *
from docx.shared import *


class DocumentTools:
    """Сборник инструментов для работы с документом"""

    def __init__(self):
        pass

    def add_content_for_cell_table(self, cell, content, text_size, alignment, bold=False, underline=False) -> None:
        """Добавить текст в ячейку таблицы с переданными параметрами"""
        text = cell.paragraphs[0].add_run(content)
        cell.paragraphs[0].alignment = alignment
        self.set_text_style_by_parameters(text, text_size, bold=bold, underline=underline)

    @staticmethod
    def add_empty_string(document) -> None:
        """Добавить пустую строку в документ"""
        document.add_paragraph("")

    @staticmethod
    def create_table_by_parameters(document, rows: int, cols: int):
        """Создать таблицу с заданным количеством строк и столбцов"""
        table = document.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.width = 1
        return table

    @staticmethod
    def set_style_standard_for_all_document(document: Document(), font_size: int, font_name: str) -> None:
        """Задать стиль для всего документа(размер и стиль шрифта)"""
        style = document.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)

    @staticmethod
    def set_text_style_by_parameters(text, text_size, bold=False, underline=False, italic=False) -> None:
        """Задать стиль текста по заданным параметрам(размер текста, жирный, подчеркнутый или курсив"""
        text.font.size = Pt(text_size)
        text.font.bold = bold
        text.font.italic = italic
        if underline:
            text.font.underline = WD_UNDERLINE.SINGLE

    @staticmethod
    def set_padding_sizes_for_all_document(document: Document(), bottom: int, top: int, right: int, left: int) -> None:
        """Добавить отступы для всего документа"""
        section = document.sections[0]
        section.bottom_margin = Mm(bottom)
        section.top_margin = Mm(top)
        section.right_margin = Mm(right)
        section.left_margin = Mm(left)
